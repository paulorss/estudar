import streamlit as st
import pandas as pd
import json
import io
import sys
import subprocess
from typing import List, Dict, Any, Optional
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import spacy
from presidio_analyzer import RecognizerRegistry, PatternRecognizer, AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

@st.cache_resource
def load_spacy_model():
    """Load Portuguese language model, downloading it if necessary"""
    try:
        # Try to load the model
        return spacy.load("pt_core_news_lg")
    except OSError:
        # If model is not found, download it
        st.info("Downloading Portuguese language model... This may take a while.")
        subprocess.check_call([
            sys.executable, 
            "-m", 
            "spacy", 
            "download", 
            "pt_core_news_lg"
        ])
        return spacy.load("pt_core_news_lg")

# Initialize spaCy
try:
    nlp = load_spacy_model()
    st.success("Modelo de linguagem carregado com sucesso!")
except Exception as e:
    st.error(f"Erro ao carregar modelo de linguagem: {str(e)}")
    st.stop()

class CustomAnalyzer:
    def __init__(self):
        """Initialize the analyzer with custom recognizers"""
        # Initialize the recognizer registry
        self.registry = RecognizerRegistry()
        
        # Add custom recognizers
        self.setup_recognizers()
        
        # Initialize the analyzer engine with custom registry
        self.analyzer = AnalyzerEngine(registry=self.registry)
        
        # Initialize the anonymizer
        self.anonymizer = AnonymizerEngine()

    def setup_recognizers(self):
        """Setup all custom recognizers for Brazilian documents and data"""
        recognizers = [
            # Documentos
            PatternRecognizer(
                supported_entity="CPF",
                patterns=[{
                    "pattern": r"\d{3}\.?\d{3}\.?\d{3}-?\d{2}",
                    "score": 0.95
                }]
            ),
            PatternRecognizer(
                supported_entity="RG",
                patterns=[{
                    "pattern": r"\d{2}\.?\d{3}\.?\d{3}[-]?\d{1}|\d{2}\.?\d{3}\.?\d{3}",
                    "score": 0.95
                }]
            ),
            PatternRecognizer(
                supported_entity="CNH",
                patterns=[{
                    "pattern": r"\d{11}",
                    "score": 0.5
                }]
            ),
            # Contato
            PatternRecognizer(
                supported_entity="EMAIL",
                patterns=[{
                    "pattern": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
                    "score": 0.95
                }]
            ),
            PatternRecognizer(
                supported_entity="PHONE",
                patterns=[{
                    "pattern": r"(?:\+55\s?)?(?:\(?\d{2}\)?[-\s]?)?\d{4,5}[-\s]?\d{4}",
                    "score": 0.95
                }]
            ),
            # Dados Bancários
            PatternRecognizer(
                supported_entity="CREDIT_CARD",
                patterns=[{
                    "pattern": r"\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}",
                    "score": 0.95
                }]
            ),
            PatternRecognizer(
                supported_entity="BANK_ACCOUNT",
                patterns=[{
                    "pattern": r"(?:ag[êe]ncia|conta)\s*:?\s*\d{1,4}[-.]?\d{1,10}",
                    "score": 0.85
                }]
            ),
            # Endereço
            PatternRecognizer(
                supported_entity="CEP",
                patterns=[{
                    "pattern": r"\d{5}[-]?\d{3}",
                    "score": 0.95
                }]
            ),
            PatternRecognizer(
                supported_entity="ADDRESS",
                patterns=[{
                    "pattern": r"(?:Rua|Av|Avenida|Alameda|Al|Praça|R|Travessa|Rod|Rodovia)\s+(?:[A-ZÀ-Ú][a-zà-ú]+\s*)+,?\s*\d+",
                    "score": 0.85
                }]
            ),
            # Nomes
            PatternRecognizer(
                supported_entity="PERSON",
                patterns=[{
                    "pattern": r"(?i)([A-ZÀ-Ú][a-zà-ú]+(?:\s+(?:dos?|das?|de|e|[A-ZÀ-Ú][a-zà-ú]+))+)",
                    "score": 0.7
                }]
            )
        ]
        
        # Add all recognizers to registry
        for recognizer in recognizers:
            self.registry.add_recognizer(recognizer)

    def analyze_text(self, text: str) -> List[Dict]:
        """Analyze text using custom recognizers"""
        if not isinstance(text, str) or not text.strip():
            return []
            
        try:
            return self.analyzer.analyze(
                text=text,
                language="pt",
                entities=[
                    "CPF", "RG", "CNH", "EMAIL", "PHONE", 
                    "CREDIT_CARD", "BANK_ACCOUNT", "CEP",
                    "ADDRESS", "PERSON"
                ]
            )
        except Exception as e:
            st.error(f"Erro ao analisar texto: {str(e)}")
            return []

    def anonymize_text(self, text: str) -> str:
        """Anonymize text using Presidio"""
        if not isinstance(text, str) or not text.strip():
            return text
            
        try:
            # Analyze text
            analyzer_results = self.analyze_text(text)
            
            # Define operators for each entity type
            operators = {
                "CPF": OperatorConfig("mask", {"chars_to_mask": 6, "masking_char": "*"}),
                "RG": OperatorConfig("mask", {"chars_to_mask": 4, "masking_char": "*"}),
                "CNH": OperatorConfig("mask", {"chars_to_mask": 6, "masking_char": "*"}),
                "EMAIL": OperatorConfig("mask", {"chars_to_mask": -4, "masking_char": "*"}),
                "PHONE": OperatorConfig("mask", {"chars_to_mask": 4, "masking_char": "*"}),
                "CREDIT_CARD": OperatorConfig("mask", {"chars_to_mask": 12, "masking_char": "*"}),
                "BANK_ACCOUNT": OperatorConfig("mask", {"chars_to_mask": -4, "masking_char": "*"}),
                "CEP": OperatorConfig("mask", {"chars_to_mask": 3, "masking_char": "*"}),
                "ADDRESS": OperatorConfig("replace", {"new_value": "[ENDEREÇO PROTEGIDO]"}),
                "PERSON": OperatorConfig("replace", {"new_value": "[NOME PROTEGIDO]"})
            }
            
            # Anonymize text
            anonymized_result = self.anonymizer.anonymize(
                text=text,
                analyzer_results=analyzer_results,
                operators=operators
            )
            
            return anonymized_result.text
            
        except Exception as e:
            st.error(f"Erro ao anonimizar texto: {str(e)}")
            return text

    def process_csv(self, content: str) -> pd.DataFrame:
        """Process CSV file"""
        try:
            df = pd.read_csv(pd.StringIO(content))
            for column in df.columns:
                if df[column].dtype == 'object':
                    df[column] = df[column].apply(lambda x: self.anonymize_text(str(x)) if pd.notna(x) else x)
            return df
        except Exception as e:
            st.error(f"Erro ao processar CSV: {str(e)}")
            return pd.DataFrame()

    def process_json(self, content: str) -> dict:
        """Process JSON file"""
        try:
            data = json.loads(content)
            
            def anonymize_dict(d: Any) -> Any:
                if isinstance(d, dict):
                    return {k: anonymize_dict(v) for k, v in d.items()}
                elif isinstance(d, list):
                    return [anonymize_dict(v) for v in d]
                elif isinstance(d, str):
                    return self.anonymize_text(d)
                else:
                    return d
                    
            return anonymize_dict(data)
        except Exception as e:
            st.error(f"Erro ao processar JSON: {str(e)}")
            return {}

    def process_pdf(self, pdf_bytes: bytes) -> bytes:
        """Process PDF file"""
        try:
            pdf_reader = PdfReader(io.BytesIO(pdf_bytes))
            pdf_writer = PdfWriter()
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    anonymized_text = self.anonymize_text(text)
                    
                    packet = io.BytesIO()
                    c = canvas.Canvas(packet, pagesize=letter)
                    
                    y = 750
                    for line in anonymized_text.split('\n'):
                        if line.strip():
                            c.drawString(50, y, line)
                            y -= 15
                    
                    c.save()
                    packet.seek(0)
                    new_page = PdfReader(packet).pages[0]
                    pdf_writer.add_page(new_page)
                else:
                    pdf_writer.add_page(page)
            
            output = io.BytesIO()
            pdf_writer.write(output)
            return output.getvalue()
            
        except Exception as e:
            st.error(f"Erro ao processar PDF: {str(e)}")
            return pdf_bytes

    def process_docx(self, docx_bytes: bytes) -> bytes:
        """Process DOCX file"""
        try:
            doc_temp = io.BytesIO(docx_bytes)
            doc = Document(doc_temp)
            anonymized_doc = Document()
            
            for para in doc.paragraphs:
                anonymized_doc.add_paragraph(self.anonymize_text(para.text))
            
            for table in doc.tables:
                new_table = anonymized_doc.add_table(
                    rows=len(table.rows),
                    cols=len(table.columns)
                )
                
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_cell = new_table.cell(i, j)
                        new_cell.text = self.anonymize_text(cell.text)
            
            output = io.BytesIO()
            anonymized_doc.save(output)
            return output.getvalue()
            
        except Exception as e:
            st.error(f"Erro ao processar DOCX: {str(e)}")
            return docx_bytes

def main():
    st.set_page_config(
        page_title="Anonimizador de Dados - LGPD",
        page_icon="🔒",
        layout="wide"
    )
    
    st.title("Anonimizador de Dados - LGPD")
    
    st.markdown("""
    ### Dados detectados e anonimizados:
    
    **Documentos Brasileiros:**
    - CPF (mascaramento parcial)
    - RG (mascaramento parcial)
    - CNH (mascaramento parcial)
    
    **Dados de Contato:**
    - E-mails (mascaramento)
    - Telefones (mascaramento)
    
    **Dados Financeiros:**
    - Cartões de Crédito (mascaramento)
    - Contas Bancárias (mascaramento)
    
    **Localização:**
    - CEP (mascaramento)
    - Endereços (substituição)
    
    **Dados Pessoais:**
    - Nomes (substituição)
    
    O sistema usa reconhecimento de padrões para detectar e anonimizar dados sensíveis em português.
    """)
    
    analyzer = CustomAnalyzer()
    
    st.header("Anonimização de Texto")
    text_input = st.text_area("Digite o texto a ser anonimizado:", height=150)
    if st.button("Anonimizar Texto") and text_input:
        anonymized_text = analyzer.anonymize_text(text_input)
        st.text_area("Texto Anonimizado:", anonymized_text, height=150)
    
    st.header("Anonimização de Arquivo")
    file = st.file_uploader(
        "Escolha um arquivo (.csv, .txt, .json, .pdf, .docx)",
        type=['csv', 'txt', 'json', 'pdf', 'docx']
    )
    
    if file is not None:
        try:
            content = file.getvalue()
            
            if file.name.endswith('.csv'):
                content_text = content.decode('utf-8')
                result = analyzer.process_csv(content_text)
                csv = result.to_csv(index=False)
                st.download_button(
                    "Download CSV Anonimizado",
                    csv,
                    "anonimizado.csv",
                    "text/csv"
                )
                st.write("Preview do arquivo anonimizado:")
                st.dataframe(result)
                
            elif file.name.endswith('.json'):
                content_text = content.decode('utf-8')
                result = analyzer.process_json(content_text)
                json_str = json.dumps(result, ensure_ascii=False, indent=2)
                st.download_button(
                    "Download JSON Anonimizado",
                    json_str,
                    "anonimizado.json",
                    "application/json"
                )
                st.write("Preview do arquivo anonimizado:")
                st.json(result)
                
            elif file.name.endswith('.docx'):
                docx_anonymized = analyzer.process_docx(content)
                st.download_button(
                    "Download DOCX Anonimizado",
                    docx_anonymized,
                    "anonimizado.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                
            elif file.name.endswith('.pdf'):
                pdf_anonymized = analyzer.process_pdf(content)
                st.download_button(
                    "Download PDF Anonimizado",
                    pdf_anonymized,
                    "anonimizado.pdf",
                    "application/pdf"
                )
                
            else:  # .txt
                content_text = content.decode('utf-8')
                anonymized_text = analyzer.anonymize_text(content_text)
                st.download_button(
                    "Download TXT Anonimizado",
                    anonymized_text,
                    "anonimizado.txt",
                    "text/plain"
                )
                st.write("Preview do arquivo anonimizado:")
                st.text_area("Conteúdo:", anonymized_text, height=300)
            
            st.success("Arquivo processado com sucesso!")
            
        except Exception as e:
            st.error(f"Erro ao processar arquivo: {str(e)}")
            
    st.markdown("""
    ### Como usar:
    1. Digite um texto diretamente na caixa de texto ou
    2. Faça upload de um arquivo suportado (.csv, .txt, .json, .pdf, .docx)
    3. Clique no botão correspondente para anonimizar
    4. Faça o download do arquivo anonimizado
    
    ### Observações:
    - Os dados são processados localmente em seu navegador
    - Nenhuma informação é armazenada ou transmitida
    - Alguns formatos complexos de PDF podem não ser processados corretamente
    """)

if __name__ == "__main__":
    main()
